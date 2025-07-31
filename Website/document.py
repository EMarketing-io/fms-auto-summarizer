from docx import Document
import io
import re


# Function to create a DOCX file in memory from a summary JSON
def create_docx_in_memory(summary_json, document_title):
    doc = Document()
    doc.add_heading(document_title, level=0)

    # Add sections based on the summary JSON
    for section in summary_json.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        
        # Add content to the section
        for line in section["content"].split("\n"):
            line = line.strip()
            
            # If the line starts with a bullet point, format it accordingly
            if line.startswith("- "):
                line = line[2:].strip()
                para = doc.add_paragraph(style="List Bullet")
                parts = re.split(r"(\*\*.*?\*\*)", line)
                
                # Add parts with bold formatting if they are enclosed in **
                for part in parts:
                    run = para.add_run()
                    if part.startswith("**") and part.endswith("**"):
                        run.text = part[2:-2]
                        run.bold = True
                    else:
                        run.text = part
            else:
                doc.add_paragraph(line.strip())

    # Save the document to a BytesIO stream
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return doc_stream